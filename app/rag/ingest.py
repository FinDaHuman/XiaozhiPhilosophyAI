"""
Knowledge Base Ingestion Pipeline

Loads philosophy documents from the data/ directory (PDF, DOCX, TXT, MD),
cleans and chunks the text, embeds with the configured local sentence-transformer,
and stores the vectors in ChromaDB.
"""

import os
import re
import sys
import unicodedata
from pathlib import Path
from typing import Optional

# Fix Windows console encoding for emoji/Unicode output
if sys.stdout and hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if sys.stderr and hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.rag.embeddings import get_embedding_model as get_embeddings
from langchain_chroma import Chroma
from dotenv import load_dotenv

load_dotenv()


# ─── Configuration ──────────────────────────────────────────────────────────

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "chroma_db")
CHROMA_COLLECTION = os.getenv("CHROMA_COLLECTION", "philosophy_docs")
DATA_DIR = os.getenv("DATA_DIR", "data")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-small")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


# ─── Document Loaders ───────────────────────────────────────────────────────

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def load_text(file_path: str) -> str:
    """Load a plain text or markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
}


# ─── Text Cleaning ──────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Clean extracted text:
    - Normalize Unicode (NFC for Vietnamese diacritics)
    - Remove excessive whitespace
    - Remove control characters
    """
    # Normalize Unicode — important for Vietnamese characters
    text = unicodedata.normalize("NFC", text)

    # Remove control characters (keep newlines and tabs)
    text = re.sub(r"[^\S \n\t]+", " ", text)

    # Collapse multiple blank lines into two
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse multiple spaces into one
    text = re.sub(r" {2,}", " ", text)

    return text.strip()


# ─── Pipeline ────────────────────────────────────────────────────────────────

def load_documents(data_dir: Optional[str] = None) -> list[Document]:
    """
    Load all supported documents from the data directory.
    Returns a list of LangChain Document objects.
    """
    data_dir = data_dir or DATA_DIR
    data_path = Path(data_dir)

    if not data_path.exists():
        raise FileNotFoundError(f"Data directory not found: {data_path.resolve()}")

    documents = []

    for file_path in sorted(data_path.iterdir()):
        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        loader = LOADERS.get(ext)
        if not loader:
            continue

        try:
            print(f"  📄 Loading: {file_path.name}")
            raw_text = loader(str(file_path))
            cleaned = clean_text(raw_text)

            if not cleaned:
                print(f"  ⚠️  Skipped (empty): {file_path.name}")
                continue

            if file_path.name.startswith("Slide") and file_path.name.endswith("_OCR.md"):
                # Split slide-OCR files whenever "## [Slide <label>]" is encountered.
                # The label may be a plain number ("12") or deck-prefixed ("KTCT 5").
                slides = re.split(r'(?=## \[Slide [^\]]+\])', cleaned)
                for slide_text in slides:
                    slide_text = slide_text.strip()
                    if not slide_text:
                        continue

                    match = re.search(r'## \[Slide ([^\]]+)\]', slide_text)
                    if match:
                        slide_label = match.group(1).strip()
                        doc = Document(
                            page_content=slide_text,
                            metadata={
                                "source": f"Slide {slide_label}",
                                "file_type": ext,
                                "file_path": str(file_path.resolve()),
                            },
                        )
                        documents.append(doc)
            else:
                doc = Document(
                    page_content=cleaned,
                    metadata={
                        "source": file_path.name,
                        "file_type": ext,
                        "file_path": str(file_path.resolve()),
                    },
                )
                documents.append(doc)
            print(f"  ✅ Loaded: {file_path.name} ({len(cleaned):,} chars)")

        except Exception as e:
            print(f"  ❌ Error loading {file_path.name}: {e}")

    return documents


def chunk_documents(documents: list[Document]) -> list[Document]:
    """Split documents into chunks for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
        length_function=len,
    )
    chunks = splitter.split_documents(documents)
    return chunks





def build_vector_store(
    chunks: list[Document],
    embeddings = None,
    persist_directory: Optional[str] = None,
) -> Chroma:
    """
    Create/overwrite the ChromaDB vector store from document chunks.
    Batches embedding requests with retry logic to respect API rate limits.
    """
    import time
    import math

    embeddings = embeddings or get_embeddings()
    persist_dir = persist_directory or CHROMA_PERSIST_DIR

    BATCH_SIZE = 500  # Larger batches for local embeddings
    MAX_RETRIES = 5
    total_batches = math.ceil(len(chunks) / BATCH_SIZE)

    print(f"\n  🗄️  Building vector store in: {persist_dir}")
    print(f"  📊 Embedding {len(chunks)} chunks in {total_batches} batches (batch size={BATCH_SIZE})...")

    vectorstore = None

    for i in range(0, len(chunks), BATCH_SIZE):
        batch_num = (i // BATCH_SIZE) + 1
        batch = chunks[i : i + BATCH_SIZE]

        # Retry loop for rate limit errors
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                print(f"\n  📦 Batch {batch_num}/{total_batches}: embedding {len(batch)} chunks...")

                if vectorstore is None:
                    vectorstore = Chroma.from_documents(
                        documents=batch,
                        embedding=embeddings,
                        collection_name=CHROMA_COLLECTION,
                        persist_directory=persist_dir,
                    )
                else:
                    vectorstore.add_documents(batch)

                print(f"  ✅ Batch {batch_num}/{total_batches} complete!")
                break  # Success — exit retry loop

            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    # Parse retry delay from error if available
                    wait_secs = 65 * attempt  # Exponential: 65s, 130s, 195s...
                    print(f"  ⚠️  Rate limited (attempt {attempt}/{MAX_RETRIES}). Waiting {wait_secs}s...")
                    time.sleep(wait_secs)
                    if attempt == MAX_RETRIES:
                        raise RuntimeError(
                            f"Failed after {MAX_RETRIES} retries due to rate limiting. "
                            f"Try again later or reduce batch size."
                        ) from e
                else:
                    raise  # Non-rate-limit error, don't retry


    return vectorstore


def load_existing_vector_store(
    embeddings = None,
    persist_directory: Optional[str] = None,
) -> Optional[Chroma]:
    """Load an existing ChromaDB vector store, if it exists."""
    embeddings = embeddings or get_embeddings()
    persist_dir = persist_directory or CHROMA_PERSIST_DIR

    if not Path(persist_dir).exists():
        return None

    return Chroma(
        collection_name=CHROMA_COLLECTION,
        embedding_function=embeddings,
        persist_directory=persist_dir,
    )


def run_ingest_pipeline(data_dir: Optional[str] = None, resume: bool = False) -> Chroma:
    """
    Full ingestion pipeline:
    1. Load documents from data/
    2. Clean text
    3. Chunk
    4. Embed
    5. Store in ChromaDB

    If resume=True, skip chunks that were already embedded in a previous run.
    """
    print("\n🚀 Starting Knowledge Base Ingestion Pipeline")
    print("=" * 50)

    # Step 1: Load
    print("\n📂 Step 1: Loading documents...")
    documents = load_documents(data_dir)
    if not documents:
        raise ValueError("No documents found to ingest!")
    print(f"\n  📊 Loaded {len(documents)} documents")

    # Step 2 & 3: Chunk (cleaning happens in load_documents)
    print("\n✂️  Step 2: Chunking documents...")
    chunks = chunk_documents(documents)
    print(f"  📊 Created {len(chunks)} chunks (size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})")

    # Check for resume
    if resume:
        existing = load_existing_vector_store()
        if existing is not None:
            try:
                existing_count = existing._collection.count()
                if existing_count > 0 and existing_count < len(chunks):
                    print(f"\n🔄 Resuming: skipping {existing_count} already-embedded chunks...")
                    chunks = chunks[existing_count:]
                    print(f"  📊 Remaining: {len(chunks)} chunks to embed")

                    # Add remaining chunks to existing store
                    print("\n🧠 Step 3: Embedding remaining chunks...")
                    vectorstore = build_vector_store_resume(chunks, existing)

                    total = existing_count + len(chunks)
                    print("\n" + "=" * 50)
                    print("✅ Knowledge Base Updated!")
                    print(f"   Previously embedded: {existing_count}")
                    print(f"   Newly embedded: {len(chunks)}")
                    print(f"   Total chunks: {total}")
                    print(f"   Store: {CHROMA_PERSIST_DIR}")
                    return vectorstore

                elif existing_count >= len(chunks):
                    print(f"\n✅ All {existing_count} chunks already embedded. Nothing to do!")
                    return existing
            except Exception as e:
                print(f"  ⚠️  Could not read existing store: {e}. Starting fresh...")

    # Step 4 & 5: Embed and Store
    print("\n🧠 Step 3: Embedding and storing...")
    vectorstore = build_vector_store(chunks)

    print("\n" + "=" * 50)
    print("✅ Knowledge Base Ready!")
    print(f"   Documents: {len(documents)}")
    print(f"   Chunks: {len(chunks)}")
    print(f"   Store: {CHROMA_PERSIST_DIR}")

    return vectorstore


def build_vector_store_resume(
    chunks: list[Document],
    existing_store: Chroma,
) -> Chroma:
    """Add new chunks to an existing vector store with rate limiting."""
    import time
    import math

    BATCH_SIZE = 20
    MAX_RETRIES = 5
    total_batches = math.ceil(len(chunks) / BATCH_SIZE)

    print(f"  📊 Embedding {len(chunks)} chunks in {total_batches} batches...")

    for i in range(0, len(chunks), BATCH_SIZE):
        batch_num = (i // BATCH_SIZE) + 1
        batch = chunks[i : i + BATCH_SIZE]

        for attempt in range(1, MAX_RETRIES + 1):
            try:
                print(f"\n  📦 Batch {batch_num}/{total_batches}: embedding {len(batch)} chunks...")
                existing_store.add_documents(batch)
                print(f"  ✅ Batch {batch_num}/{total_batches} complete!")
                break
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    wait_secs = 65 * attempt
                    print(f"  ⚠️  Rate limited (attempt {attempt}/{MAX_RETRIES}). Waiting {wait_secs}s...")
                    time.sleep(wait_secs)
                    if attempt == MAX_RETRIES:
                        raise RuntimeError(
                            f"Failed after {MAX_RETRIES} retries. "
                            f"Embedded up to batch {batch_num-1}. Run with --resume again later."
                        ) from e
                else:
                    raise

        if i + BATCH_SIZE < len(chunks):
            print(f"  ⏳ Cooling down 65s before next batch...")
            time.sleep(65)

    print(f"\n  ✅ All remaining chunks embedded!")
    return existing_store


# ─── CLI Entry Point ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    resume = "--resume" in sys.argv
    run_ingest_pipeline(resume=resume)

